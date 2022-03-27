# install instructions
# pip install python-docx

from docx import Document
import csv


def add_schedule_page(document, schedule, number_of_rounds):
    document.add_heading(schedule['name'], level=1)
    document.add_paragraph('table: ' +  schedule['table'])

    i = 1 
    for round in schedule['rounds'][0:number_of_rounds]:        
        document.add_paragraph(round,style='List Number')
        i += 1 
    document.add_page_break()


schedules = []
file_path = '/home/jackson/source/speed-date-backend/stals-speeddating-backend/printed_schedule_generator/dateschedule_example.csv'
max_round_index = 0
with open(file_path, newline='') as csvfile:
    spamreader = csv.reader(csvfile, delimiter=',', quotechar=',')
    for row in spamreader:
        if len(row) == 0: continue
        if(row[0] == "Women Schedule" or row[0] == "Men Schedule"):
            continue
        person_schedule = {
            'name': row[0],
            'table': row[2],
            'rounds': []
        }
        i = 0
        for word in row[6:]:
            if word.strip() != 'Break' and len(word) > 0 and i > max_round_index: 
                max_round_index += 1
            person_schedule['rounds'].append(word)
            i += 1

        schedules.append(person_schedule)


document = Document()

for schedule in schedules:
    add_schedule_page(document, schedule, max_round_index)

document.save('schedule.docx')
